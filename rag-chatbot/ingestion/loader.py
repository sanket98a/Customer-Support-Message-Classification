from __future__ import annotations

from pathlib import Path
from typing import Sequence

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document

from config.config import AppConfig
from utils.helper import clean_text, ensure_directory, is_supported_file
from utils.logger import get_logger


class DocumentLoader:
    """Load supported documents from disk into LangChain documents."""

    def __init__(self, raw_dir: str | Path | None = None, config: AppConfig | None = None) -> None:
        self.config = config or AppConfig()
        self.raw_dir = Path(raw_dir or self.config.raw_dir)
        ensure_directory(self.raw_dir)
        self.logger = get_logger("DocumentLoader")

    def load_documents(self, paths: Sequence[str | Path] | None = None) -> list[Document]:
        """Load all supported files from the supplied paths or the raw directory."""
        if paths is None:
            file_paths = sorted(self.raw_dir.iterdir())
        else:
            file_paths = [Path(path) for path in paths]

        documents: list[Document] = []
        for file_path in file_paths:
            if not file_path.is_file():
                continue
            if not is_supported_file(file_path, self.config.supported_extensions):
                self.logger.warning("Unsupported file type ignored: %s", file_path)
                continue

            text = self._extract_text(file_path)
            if not text.strip():
                self.logger.warning("No text found in document: %s", file_path)
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "file_path": str(file_path),
                        "document_name": file_path.stem,
                        "file_type": file_path.suffix.lower(),
                    },
                )
            )

        return documents

    def _extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._read_pdf(file_path)
        if suffix == ".docx":
            return self._read_docx(file_path)
        return self._read_text_file(file_path)

    def _read_pdf(self, file_path: Path) -> str:
        text_parts: list[str] = []
        with fitz.open(file_path) as document:
            for page_number, page in enumerate(document, start=1):
                page_text = page.get_text()
                if page_text:
                    text_parts.append(f"[Page {page_number}]\n{page_text}")
        return clean_text("\n\n".join(text_parts))

    def _read_docx(self, file_path: Path) -> str:
        document = DocxDocument(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return clean_text("\n".join(paragraphs))

    def _read_text_file(self, file_path: Path) -> str:
        return clean_text(file_path.read_text(encoding="utf-8", errors="ignore"))
